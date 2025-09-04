"""
Create a sample Word document for testing
This script creates a sample Chinese Word document for testing the conversion pipeline
"""

try:
    import docx
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from pathlib import Path
    
    def create_sample_word_document():
        """Create a sample Word document with Chinese content"""
        
        # Create a new document
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading('第一章：新的开始', level=1)
        
        # Add subtitle
        subtitle = doc.add_heading('小明的村庄', level=2)
        
        # Add paragraph 1
        p1 = doc.add_paragraph(
            '这是一个美丽的小村庄，坐落在青山绿水之间。村庄不大，大约有一百多户人家，'
            '每家每户都有自己的小院子，院子里种着各种蔬菜和花草。'
        )
        
        # Add paragraph 2
        p2 = doc.add_paragraph(
            '小明是这个村庄里最活泼的孩子之一。他今年十二岁，有着一双明亮的眼睛和灿烂的笑容。'
            '从小，小明就对外面的世界充满了好奇。'
        )
        
        # Add another heading
        heading2 = doc.add_heading('日常生活', level=2)
        
        # Add a subheading
        subheading = doc.add_heading('早晨', level=3)
        
        # Add a paragraph with some formatting
        p3 = doc.add_paragraph('每天早晨，小明都会早早起床，帮助父母干一些力所能及的家务活。他会：')
        
        # Add a list
        doc.add_paragraph('喂鸡喂鸭', style='List Bullet')
        doc.add_paragraph('浇菜园里的蔬菜', style='List Bullet')  
        doc.add_paragraph('帮母亲准备早餐', style='List Bullet')
        
        # Add another subheading
        subheading2 = doc.add_heading('上学路上', level=3)
        
        # Add numbered list
        doc.add_paragraph('上学的路上，小明总是和好朋友小红一起走。他们会经过：')
        doc.add_paragraph('村口的老槐树', style='List Number')
        doc.add_paragraph('清澈的小溪', style='List Number')
        doc.add_paragraph('开满野花的山坡', style='List Number')
        
        # Add paragraph with bold text
        p4 = doc.add_paragraph()
        p4.add_run('小明经常对小红说："')
        bold_run = p4.add_run('我想去看看山那边是什么样子的。')
        bold_run.bold = True
        p4.add_run('"')
        
        # Add table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '时间'
        hdr_cells[1].text = '活动'
        hdr_cells[2].text = '地点'
        
        # Add data rows
        schedule_data = [
            ('6:00-7:00', '做家务', '家中'),
            ('7:00-8:00', '吃早餐', '家中'),
            ('8:00-12:00', '上课', '学校'),
            ('14:00-17:00', '上课', '学校'),
            ('17:00-19:00', '玩耍', '村庄'),
            ('19:00-21:00', '做作业', '家中')
        ]
        
        for time, activity, location in schedule_data:
            row_cells = table.add_row().cells
            row_cells[0].text = time
            row_cells[1].text = activity
            row_cells[2].text = location
        
        # Add final paragraph
        doc.add_paragraph(
            '这就是小明平凡而充实的一天。但是在这平凡的生活中，'
            '他的心里已经种下了梦想的种子。'
        )
        
        # Save the document
        output_dir = Path('./input_documents')
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / 'sample_chapter.docx'
        
        doc.save(str(output_path))
        
        print(f"✅ Sample Word document created: {output_path}")
        return str(output_path)
    
    if __name__ == "__main__":
        create_sample_word_document()

except ImportError:
    print("⚠️ python-docx not installed. Install requirements first:")
    print("pip install python-docx")
except Exception as e:
    print(f"❌ Error creating sample document: {e}")
